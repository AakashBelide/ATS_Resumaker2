"""Generate candidate files to upload into OpenCATS for the ranking test.

Writes our real resume PDF + one PDF per decoy into ./candidates/ (gitignored -
our resume carries PII). Run from the repo's `resumaker/` env:

    cd resumaker && uv run python ../validation/opencats/make_candidates.py
"""
from __future__ import annotations

import shutil
import sys
from pathlib import Path

# make the resumaker package importable regardless of cwd
sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "resumaker"))

from docx import Document  # noqa: E402

from pocs.ats_sim.decoys import DECOYS  # noqa: E402
from pocs.resume.render_pdf import docx_to_pdf  # noqa: E402

_HERE = Path(__file__).resolve().parent
_OUT = _HERE / "candidates"
_RESUME = (_HERE.parents[1] / "outputs" / "state-street-ai-orchestration-engineer" /
           "state-street-ai-orchestration-engineer.pdf")


def _decoy_pdf(label: str, text: str, dest: Path) -> None:
    doc = Document()
    doc.add_heading(label.replace("decoy_", "").replace("_", " ").title(), level=1)
    for para in text.split(". "):
        if para.strip():
            doc.add_paragraph(para.strip().rstrip(".") + ".")
    docx = dest.with_suffix(".docx")
    doc.save(str(docx))
    docx_to_pdf(str(docx))
    docx.unlink(missing_ok=True)


def main() -> None:
    _OUT.mkdir(parents=True, exist_ok=True)
    if _RESUME.exists():
        shutil.copy(_RESUME, _OUT / "00_OURS_aakash_belide.pdf")
        print(f"copied our resume -> {_OUT / '00_OURS_aakash_belide.pdf'}")
    else:
        print(f"WARNING: our resume not found at {_RESUME}")
    for i, (label, text) in enumerate(DECOYS, 1):
        dest = _OUT / f"{i:02d}_{label}.pdf"
        _decoy_pdf(label, text, dest)
        print(f"wrote decoy -> {dest}")
    print(f"\n{len(list(_OUT.glob('*.pdf')))} candidate PDFs in {_OUT}")


if __name__ == "__main__":
    main()
