"""Render a ResumeContent to an ATS-safe .docx (Task 1.8 render step).

ATS-safe choices (blueprint §4, §10): single column, standard font, standard
section headers as a bold paragraph with a bottom border (NOT a table), dates on a
right-aligned tab stop (same logical line, no columns), real w:hyperlink relations,
contact info in the BODY (not header/footer), US Letter, no images/tables/text-boxes.
`**bold**` markers in bullets become real bold runs (metric/keyword emphasis).
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from core import profile as prof
from core.schemas import ResumeContent

FONT = "Calibri"
BODY_PT = 10.5
RIGHT_TAB = Inches(7.5)   # US Letter 8.5in - 2*0.5in margins

# ATS-safe ASCII normalization (blueprint §4/§10, career-ops lesson): non-ASCII
# glyphs corrupt PDF text extraction AND em-dashes read as an AI tell to recruiters.
_ASCII_MAP = {
    "—": "-", "–": "-", "‒": "-", "‐": "-", "‑": "-",
    "‘": "'", "’": "'", "“": '"', "”": '"',
    "…": "...", "•": "-", "·": "-", "→": " to ", "←": " ",
    " ": " ", "​": "", "﻿": "", "‑": "-",
}


def _ascii(text: str) -> str:
    if not text:
        return text
    for bad, good in _ASCII_MAP.items():
        text = text.replace(bad, good)
    # "->" reads awkwardly in prose; make it "to" (keeps it ASCII + readable)
    text = re.sub(r"\s*->\s*", " to ", text)
    return text


def _set_margins(section):
    section.page_width = Inches(8.5)      # US Letter
    section.page_height = Inches(11.0)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(0.5))


def _base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(BODY_PT)
    pf = st.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0


def _add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0000FF"); rpr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _section_header(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(BODY_PT + 0.5)
    # bottom border on the paragraph (native heading rule, not a table)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "auto")):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def _add_markdown(paragraph, text):
    """Render **bold** spans as real bold runs; rest plain. ASCII-normalized."""
    text = _ascii(text)
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if not part:
            continue
        run = paragraph.add_run(part[2:-2] if part.startswith("**") else part)
        if part.startswith("**"):
            run.bold = True


def _heading_row(doc, left_bold, left_rest, right):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(RIGHT_TAB, alignment=WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(left_bold); r.bold = True
    if left_rest:
        p.add_run(left_rest)
    if right:
        rr = p.add_run(f"\t{right}"); rr.bold = True
    return p


def _as_text(b) -> str:
    """Coerce a bullet to a string (LLMs sometimes emit {'text': ...} objects)."""
    if isinstance(b, dict):
        return str(b.get("text") or b.get("bullet") or b.get("content") or "")
    return str(b)


def _bullets(doc, items):
    for b in items:
        text = _as_text(b)
        if not text.strip():
            continue
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.left_indent = Inches(0.25)
        bp.paragraph_format.first_line_indent = Inches(-0.25)
        bp.paragraph_format.space_after = Pt(1)
        _add_markdown(bp, text)


def render_docx(content: ResumeContent, out_path: str,
                contact: dict | None = None, links: dict | None = None,
                include_certs: bool = False) -> str:
    p = prof.load_profile()
    contact = contact or p["contact"]
    links = links or p["links"]

    doc = Document()
    _set_margins(doc.sections[0])
    _base_style(doc)

    # --- Header (name + optional headline + contact line, all in body) ---
    name_p = doc.add_paragraph()
    name_p.alignment = 1  # center
    nr = name_p.add_run(contact["name"]); nr.bold = True; nr.font.size = Pt(17)

    if content.headline:
        hp = doc.add_paragraph(); hp.alignment = 1
        hp.add_run(_ascii(content.headline)).font.size = Pt(BODY_PT + 0.5)

    contact_p = doc.add_paragraph(); contact_p.alignment = 1
    bits = [contact.get("location", ""), contact.get("phone", ""), contact.get("email", "")]
    contact_p.add_run(" | ".join(b for b in bits if b))
    # links as real hyperlinks (portfolio + linkedin + one github per policy)
    for label, key in (("Portfolio", "portfolio"), ("LinkedIn", "linkedin"),
                       ("GitHub", "github_academic")):
        if links.get(key):
            contact_p.add_run(" | ")
            _add_hyperlink(contact_p, links[key], label)

    # --- Summary ---
    if content.summary:
        _section_header(doc, "Summary")
        sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(2)
        _add_markdown(sp, content.summary)

    # --- Skills ---
    if content.skills:
        _section_header(doc, "Skills")
        for cat, items in content.skills.items():
            if not items:
                continue
            sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(1)
            r = sp.add_run(_ascii(f"{cat}: ")); r.bold = True
            sp.add_run(_ascii(" | ".join(items)))

    # --- Experience: one clean line "Company - Title | Location .... Dates".
    #     Concise (JD-aware) titles keep this on a single line. ---
    if content.experiences:
        _section_header(doc, "Experience")
        for e in content.experiences:
            org = _ascii(e.get("organization", ""))
            title = _ascii(e.get("title", ""))
            loc = _ascii(e.get("location", ""))
            left_bold = " - ".join(x for x in (org, title) if x)
            left_rest = f"  |  {loc}" if loc else ""
            _heading_row(doc, left_bold, left_rest, e.get("dates", ""))
            _bullets(doc, e.get("bullets", []))

    # --- Projects (title rendered as a hyperlink when a url is present) ---
    if content.projects:
        _section_header(doc, "Projects")
        for pr in content.projects:
            p_title = _ascii(pr.get("title", ""))
            url = pr.get("url") or ""
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.tab_stops.add_tab_stop(
                RIGHT_TAB, alignment=WD_TAB_ALIGNMENT.RIGHT)
            if url:
                _add_hyperlink(para, url, p_title)
            else:
                para.add_run(p_title).bold = True
            if pr.get("dates"):
                para.add_run(f"\t{pr['dates']}").bold = True
            _bullets(doc, pr.get("bullets", []))

    # --- Education ---
    edu = content.education or p.get("education", [])
    if edu:
        _section_header(doc, "Education")
        for ed in edu:
            _heading_row(doc, _ascii(ed.get("organization", "")),
                         f"  |  {_ascii(ed.get('title',''))}", ed.get("dates", ""))
            extra = []
            if ed.get("gpa"):
                extra.append(f"GPA: {ed['gpa']}")
            if extra:
                ep = doc.add_paragraph(); ep.add_run(" | ".join(extra)).italic = True

    # --- Certifications (off by default: low-signal for general AI/eng roles;
    #     the space is better used for impact bullets) ---
    certs = content.certifications or p.get("certifications", [])
    if include_certs and certs:
        _section_header(doc, "Certifications")
        cp = doc.add_paragraph()
        cp.add_run(_ascii(" | ".join(c.get("title", "") for c in certs if c.get("title"))))

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
